import os

# ---- Environment setup ----
# Hugging Face Spaces writable dirs (both ephemeral here)
TMP_DIR = "/tmp/.mem0"
PERSIST_DIR = "/tmp/.embedchain"

# Override embedchain + mem0 default paths
os.environ["EMBEDCHAIN_CONFIG_DIR"] = PERSIST_DIR
os.environ["MEM0_DIR"] = TMP_DIR
os.environ["CREWAI_STORAGE_PATH"] = "/tmp/crewai_storage"
os.environ["XDG_DATA_HOME"] = os.environ["CREWAI_STORAGE_PATH"]
os.environ["HOME"] = "/tmp"
# Make sure dirs exist
os.makedirs(PERSIST_DIR, exist_ok=True)
os.makedirs(TMP_DIR, exist_ok=True)
os.makedirs(os.environ["CREWAI_STORAGE_PATH"], exist_ok=True)
os.makedirs(os.path.join("/tmp", ".local", "share"), exist_ok=True) 
from crewai.tools import BaseTool
import os, json, subprocess, sys, base64, dill
from typing import List, Dict, Any
from pydantic import Field
from dotenv import load_dotenv
import pandas as pd
import csv
import docx
import PyPDF2
import zipfile
from pathlib import Path

load_dotenv()

# ---- File & Directory Tools (keep the basic ones) ----
from crewai_tools import (
    DirectoryReadTool,
    FileReadTool
)

# ---- Web / Scraping Tools ----
from crewai_tools import (
    ScrapeWebsiteTool,
    ScrapeElementFromWebsiteTool,
    FirecrawlSearchTool,
    FirecrawlCrawlWebsiteTool,
    FirecrawlScrapeWebsiteTool,
    WebsiteSearchTool,
    ApifyActorsTool,
    BrowserbaseLoadTool,
)

# ---- Image / Vision Tools ----
from crewai_tools import VisionTool

# ---- Code / Analysis Tools ----
from crewai_tools import CodeInterpreterTool, ComposioTool

# ---- YouTube Tools ----


# ---- GitHub Tools ----


# ---- Base64 persistence ----
def persist_base64_files(base64_files: Dict[str, str], workdir: str) -> Dict[str, str]:
    saved = {}
    for fname, b64content in base64_files.items():
        path = os.path.join(workdir, os.path.basename(fname))
        with open(path, "wb") as f:
            f.write(base64.b64decode(b64content))
        saved[fname] = path
    return saved

# ---- File uploads persistence ----
def persist_uploads(files: List, workdir: str) -> Dict[str, str]:
    saved = {}
    for f in files:
        path = os.path.join(workdir, os.path.basename(f.filename))
        with open(path, "wb") as out:
            out.write(f.file.read())
        saved[f.filename] = path
    return saved

# ---- Custom File Processing Tools ----

class CSVProcessorTool(BaseTool):
    name: str = "csv_processor"
    description: str = "Read and process CSV files. Can extract specific columns, filter data, or perform basic analysis."
    workdir: str
    saved_files: Dict[str, str] = Field(default_factory=dict)

    def _run(self, filename: str, operation: str = "read", **kwargs) -> str:
        """
        Operations: 'read', 'head', 'columns', 'filter', 'stats'
        kwargs: columns, conditions, limit, etc.
        """
        try:
            if filename in self.saved_files:
                file_path = self.saved_files[filename]
            else:
                file_path = os.path.join(self.workdir, filename)
            
            if not os.path.exists(file_path):
                return json.dumps({"error": f"File {filename} not found"})
            
            df = pd.read_csv(file_path)
            
            if operation == "read":
                limit = kwargs.get("limit", 100)
                return json.dumps({"data": df.head(limit).to_dict("records"), "shape": df.shape})
            elif operation == "head":
                n = kwargs.get("n", 5)
                return json.dumps({"head": df.head(n).to_dict("records"), "shape": df.shape})
            elif operation == "columns":
                return json.dumps({"columns": list(df.columns), "dtypes": df.dtypes.astype(str).to_dict()})
            elif operation == "stats":
                return json.dumps({"stats": df.describe().to_dict()})
            elif operation == "filter":
                # Simple filtering by column values
                col = kwargs.get("column")
                value = kwargs.get("value")
                if col and col in df.columns:
                    filtered = df[df[col] == value]
                    return json.dumps({"filtered_data": filtered.to_dict("records"), "count": len(filtered)})
                return json.dumps({"error": "Column not specified or not found"})
            else:
                return json.dumps({"error": f"Unknown operation: {operation}"})
                
        except Exception as e:
            return json.dumps({"error": f"CSV processing error: {str(e)}"})

class JSONProcessorTool(BaseTool):
    name: str = "json_processor"
    description: str = "Read and process JSON files. Can extract specific keys, filter data, or analyze structure."
    workdir: str
    saved_files: Dict[str, str] = Field(default_factory=dict)

    def _run(self, filename: str, operation: str = "read", **kwargs) -> str:
        """
        Operations: 'read', 'keys', 'extract', 'search'
        """
        try:
            if filename in self.saved_files:
                file_path = self.saved_files[filename]
            else:
                file_path = os.path.join(self.workdir, filename)
            
            if not os.path.exists(file_path):
                return json.dumps({"error": f"File {filename} not found"})
            
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            if operation == "read":
                return json.dumps({"data": data, "type": type(data).__name__})
            elif operation == "keys":
                if isinstance(data, dict):
                    return json.dumps({"keys": list(data.keys())})
                elif isinstance(data, list) and data and isinstance(data[0], dict):
                    return json.dumps({"keys": list(data[0].keys())})
                else:
                    return json.dumps({"error": "Data is not a dict or list of dicts"})
            elif operation == "extract":
                key = kwargs.get("key")
                if isinstance(data, dict) and key in data:
                    return json.dumps({"extracted": data[key]})
                return json.dumps({"error": f"Key {key} not found"})
            else:
                return json.dumps({"error": f"Unknown operation: {operation}"})
                
        except Exception as e:
            return json.dumps({"error": f"JSON processing error: {str(e)}"})

class TextProcessorTool(BaseTool):
    name: str = "text_processor"
    description: str = "Read and process text files. Can search for patterns, extract lines, or analyze content."
    workdir: str
    saved_files: Dict[str, str] = Field(default_factory=dict)

    def _run(self, filename: str, operation: str = "read", **kwargs) -> str:
        """
        Operations: 'read', 'search', 'lines', 'word_count'
        """
        try:
            if filename in self.saved_files:
                file_path = self.saved_files[filename]
            else:
                file_path = os.path.join(self.workdir, filename)
            
            if not os.path.exists(file_path):
                return json.dumps({"error": f"File {filename} not found"})
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if operation == "read":
                limit = kwargs.get("limit", 5000)  # Character limit for display
                return json.dumps({"content": content[:limit], "total_chars": len(content)})
            elif operation == "search":
                pattern = kwargs.get("pattern", "")
                lines = content.split('\n')
                matches = [line for line in lines if pattern.lower() in line.lower()]
                return json.dumps({"matches": matches[:50], "total_matches": len(matches)})
            elif operation == "lines":
                start = kwargs.get("start", 0)
                end = kwargs.get("end", 10)
                lines = content.split('\n')[start:end]
                return json.dumps({"lines": lines})
            elif operation == "word_count":
                words = len(content.split())
                lines = len(content.split('\n'))
                chars = len(content)
                return json.dumps({"words": words, "lines": lines, "characters": chars})
            else:
                return json.dumps({"error": f"Unknown operation: {operation}"})
                
        except Exception as e:
            return json.dumps({"error": f"Text processing error: {str(e)}"})

class DOCXProcessorTool(BaseTool):
    name: str = "docx_processor"
    description: str = "Read and process DOCX files. Can extract text, search content, or analyze structure."
    workdir: str
    saved_files: Dict[str, str] = Field(default_factory=dict)

    def _run(self, filename: str, operation: str = "read", **kwargs) -> str:
        """
        Operations: 'read', 'search', 'paragraphs'
        """
        try:
            if filename in self.saved_files:
                file_path = self.saved_files[filename]
            else:
                file_path = os.path.join(self.workdir, filename)
            
            if not os.path.exists(file_path):
                return json.dumps({"error": f"File {filename} not found"})
            
            doc = docx.Document(file_path)
            
            if operation == "read":
                text = '\n'.join([para.text for para in doc.paragraphs])
                limit = kwargs.get("limit", 5000)
                return json.dumps({"content": text[:limit], "total_chars": len(text), "paragraphs": len(doc.paragraphs)})
            elif operation == "search":
                pattern = kwargs.get("pattern", "")
                matches = []
                for i, para in enumerate(doc.paragraphs):
                    if pattern.lower() in para.text.lower():
                        matches.append({"paragraph": i, "text": para.text})
                return json.dumps({"matches": matches[:20]})
            elif operation == "paragraphs":
                start = kwargs.get("start", 0)
                end = kwargs.get("end", 5)
                paras = [doc.paragraphs[i].text for i in range(start, min(end, len(doc.paragraphs)))]
                return json.dumps({"paragraphs": paras})
            else:
                return json.dumps({"error": f"Unknown operation: {operation}"})
                
        except Exception as e:
            return json.dumps({"error": f"DOCX processing error: {str(e)}"})

import os, json
from pathlib import Path
from crewai.tools import BaseTool
import pdfplumber
import easyocr

class AdvancedPDFProcessorTool(BaseTool):
    name: str = "advanced_pdf_processor"
    description: str = (
        "Extract text from PDF files, including scanned PDFs and images. "
        "Supports operations: 'read', 'pages', 'page_count', 'search'."
    )
    workdir: str
    saved_files: dict = {}

    def __post_init__(self):
        # Initialize EasyOCR reader once
        self.reader = easyocr.Reader(['en'], gpu=False)

    def _run(self, filename: str, operation: str = "read", **kwargs) -> str:
        try:
            # Get file path
            file_path = self.saved_files.get(filename, os.path.join(self.workdir, filename))
            if not os.path.exists(file_path):
                return json.dumps({"error": f"File {filename} not found"})

            # Open PDF with pdfplumber
            text_pages = []
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    # If no text found, try OCR
                    if not page_text or len(page_text.strip()) < 5:
                        # Convert page to image and OCR
                        image = page.to_image(resolution=300).original
                        ocr_result = self.reader.readtext(image)
                        page_text = " ".join([t[1] for t in ocr_result])
                    text_pages.append(page_text)

            # Operations
            if operation == "read":
                limit = kwargs.get("limit", 10000)
                full_text = "\n".join(text_pages)
                return json.dumps({
                    "content": full_text[:limit],
                    "total_chars": len(full_text),
                    "pages": len(text_pages)
                })

            elif operation == "pages":
                start = kwargs.get("start", 0)
                end = kwargs.get("end", len(text_pages))
                return json.dumps({"pages": text_pages[start:end]})

            elif operation == "page_count":
                return json.dumps({"page_count": len(text_pages)})

            elif operation == "search":
                pattern = kwargs.get("pattern", "").lower()
                matches = []
                for i, page_text in enumerate(text_pages):
                    if pattern in page_text.lower():
                        matches.append({"page": i+1, "snippet": page_text[:500]})
                return json.dumps({"matches": matches, "total_matches": len(matches)})

            else:
                return json.dumps({"error": f"Unknown operation: {operation}"})

        except Exception as e:
            return json.dumps({"error": f"PDF processing error: {str(e)}"})

class ExcelProcessorTool(BaseTool):
    name: str = "excel_processor"
    description: str = "Read and process Excel files (XLSX, XLS). Can read sheets, extract data, or analyze structure."
    workdir: str
    saved_files: Dict[str, str] = Field(default_factory=dict)

    def _run(self, filename: str, operation: str = "read", **kwargs) -> str:
        """
        Operations: 'read', 'sheets', 'sheet_data', 'columns'
        """
        try:
            if filename in self.saved_files:
                file_path = self.saved_files[filename]
            else:
                file_path = os.path.join(self.workdir, filename)
            
            if not os.path.exists(file_path):
                return json.dumps({"error": f"File {filename} not found"})
            
            if operation == "sheets":
                xl_file = pd.ExcelFile(file_path)
                return json.dumps({"sheet_names": xl_file.sheet_names})
            elif operation == "read" or operation == "sheet_data":
                sheet_name = kwargs.get("sheet_name", 0)  # Default to first sheet
                limit = kwargs.get("limit", 100)
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                return json.dumps({"data": df.head(limit).to_dict("records"), "shape": df.shape, "columns": list(df.columns)})
            elif operation == "columns":
                sheet_name = kwargs.get("sheet_name", 0)
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                return json.dumps({"columns": list(df.columns), "dtypes": df.dtypes.astype(str).to_dict()})
            else:
                return json.dumps({"error": f"Unknown operation: {operation}"})
                
        except Exception as e:
            return json.dumps({"error": f"Excel processing error: {str(e)}"})

class FileAnalyzerTool(BaseTool):
    name: str = "file_analyzer"
    description: str = "Analyze uploaded files to determine type, size, and basic metadata."
    workdir: str
    saved_files: Dict[str, str] = Field(default_factory=dict)

    def _run(self, filename: str = None) -> str:
        """
        Analyze a specific file or all saved files if no filename provided
        """
        try:
            if filename:
                files_to_analyze = {filename: self.saved_files.get(filename, os.path.join(self.workdir, filename))}
            else:
                files_to_analyze = self.saved_files
            
            analysis = {}
            for fname, fpath in files_to_analyze.items():
                if not os.path.exists(fpath):
                    analysis[fname] = {"error": "File not found"}
                    continue
                
                stat = os.stat(fpath)
                file_info = {
                    "size_bytes": stat.st_size,
                    "extension": Path(fpath).suffix.lower(),
                    "exists": True
                }
                
                # Determine file type
                ext = file_info["extension"]
                if ext in ['.csv']:
                    file_info["type"] = "csv"
                    try:
                        df = pd.read_csv(fpath)
                        file_info["rows"] = len(df)
                        file_info["columns"] = len(df.columns)
                        file_info["column_names"] = list(df.columns)
                    except:
                        file_info["type"] = "csv_error"
                elif ext in ['.xlsx', '.xls']:
                    file_info["type"] = "excel"
                    try:
                        xl_file = pd.ExcelFile(fpath)
                        file_info["sheets"] = xl_file.sheet_names
                    except:
                        file_info["type"] = "excel_error"
                elif ext in ['.json']:
                    file_info["type"] = "json"
                elif ext in ['.txt', '.md']:
                    file_info["type"] = "text"
                    with open(fpath, 'r', encoding='utf-8') as f:
                        content = f.read()
                        file_info["lines"] = len(content.split('\n'))
                        file_info["words"] = len(content.split())
                elif ext in ['.docx']:
                    file_info["type"] = "docx"
                elif ext in ['.pdf']:
                    file_info["type"] = "pdf"
                else:
                    file_info["type"] = "unknown"
                
                analysis[fname] = file_info
            
            return json.dumps({"file_analysis": analysis})
                
        except Exception as e:
            return json.dumps({"error": f"File analysis error: {str(e)}"})

# ---- Python Execution Tool ----
class PythonExecTool(BaseTool):
    name: str = "python_exec"
    description: str = "Run Python code in isolated subprocess with timeout and resource limits."
    workdir: str
    saved_files: Dict[str, str] = Field(default_factory=dict)
    timeout_sec: int = 60

    def __post_init__(self):
        os.makedirs(self.workdir, exist_ok=True)

    def _run(self, code_list_json: str) -> str:
        try:
            code_list = json.loads(code_list_json)
        except Exception as e:
            return json.dumps({"error": f"Invalid JSON: {e}"})

        results = []
        for item in code_list:
            step_id = item.get("step_id")
            filename = os.path.basename(item.get("filename", f"step_{step_id or 'x'}.py"))
            code = item.get("code", "")
            for name, path in self.saved_files.items():
                code = code.replace(f'"{name}"', f'"{path}"').replace(f"'{name}'", f"'{path}'")
            fpath = os.path.join(self.workdir, filename)
            with open(fpath, "w", encoding="utf-8") as fp:
                fp.write(code)
            res = self._run_file(fpath)
            res.update({"step_id": step_id, "filename": filename})
            results.append(res)
        return json.dumps(results)

    def _run_file(self, path: str) -> Dict[str, Any]:
        state_file = os.path.join(self.workdir, "_state.pkl")
        wrapper = f"""
import dill, os
if os.path.exists(r"{state_file}"):
    with open(r"{state_file}", "rb") as f:
        globals().update(dill.load(f))
with open(r"{path}", "r", encoding="utf-8") as f:
    code = f.read()
exec(code, globals())
with open(r"{state_file}", "wb") as f:
    dill.dump(globals(), f)
"""
        temp_path = path + ".wrapped.py"
        with open(temp_path, "w", encoding="utf-8") as f:
            f.write(wrapper)
        process = subprocess.Popen(
            [sys.executable, temp_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )
        stdout, stderr = process.communicate(timeout=self.timeout_sec)
        return {"returncode": process.returncode, "stdout": stdout, "stderr": stderr}

# ---- Build all tools ----
def build_all_tools(workdir: str, saved_files: Dict[str, str]):
    tools = []

    # ---- File / Directory Navigation ----
    tools += [DirectoryReadTool(), FileReadTool()]

    # ---- File Processing Tools ----
    tools += [
        CSVProcessorTool(workdir=workdir, saved_files=saved_files),
        ExcelProcessorTool(workdir=workdir, saved_files=saved_files),
        JSONProcessorTool(workdir=workdir, saved_files=saved_files),
        TextProcessorTool(workdir=workdir, saved_files=saved_files),
        DOCXProcessorTool(workdir=workdir, saved_files=saved_files),
        AdvancedPDFProcessorTool(workdir=workdir, saved_files=saved_files),
        FileAnalyzerTool(workdir=workdir, saved_files=saved_files),
    ]

    # ---- Python Execution Tool ----
    tools += [PythonExecTool(workdir=workdir, saved_files=saved_files)]

    # ---- Web Tools ----
    tools += [
        ScrapeWebsiteTool()
    ]

    return tools
